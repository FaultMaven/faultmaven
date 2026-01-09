"""ingestion.py

Purpose: Knowledge base population pipeline

Requirements:
--------------------------------------------------------------------------------
• Handle background processing
• Support multiple file formats
• Generate embeddings and store via IVectorBackend interface

Key Components:
--------------------------------------------------------------------------------
  class KnowledgeIngester: ...
  @background_task def ingest_document(...)

Technology Stack:
--------------------------------------------------------------------------------
PyPDF2, python-docx, sentence-transformers

Core Design Principles:
--------------------------------------------------------------------------------
• Privacy-First: Sanitize all external-bound data
• Resilience: Implement retries and fallbacks
• Cost-Efficiency: Use semantic caching
• Extensibility: Use interfaces for pluggable components
• Observability: Add tracing spans for key operations
• Deployment Agnostic: Use IVectorBackend interface (no direct ChromaDB imports)
"""

import logging
import os
import uuid
from typing import Any, Dict, List, Optional

import pandas as pd
import pypdf
from docx import Document

from faultmaven.models import KnowledgeBaseDocument
from faultmaven.infrastructure.observability.tracing import trace
from faultmaven.infrastructure.security.redaction import DataSanitizer
from faultmaven.infrastructure.model_cache import model_cache
from faultmaven.infrastructure.vector.base import IVectorBackend, VectorDocument, VectorSearchResult


class KnowledgeIngester:
    """Handles asynchronous ingestion of documents into the knowledge base.

    This class is deployment-agnostic and uses the IVectorBackend interface
    for vector storage operations. The concrete backend (ChromaDB, Pinecone, etc.)
    is injected via constructor dependency injection.

    Principle 1 Compliance: No direct chromadb imports - uses IVectorBackend.
    """

    def __init__(
        self,
        vector_backend: Optional[IVectorBackend] = None,
        collection_name: str = "faultmaven_kb",
        settings=None
    ):
        """Initialize KnowledgeIngester with vector backend.

        Args:
            vector_backend: IVectorBackend implementation. If None, uses factory.
            collection_name: Name of the collection for document storage.
            settings: Optional settings instance.
        """
        self.logger = logging.getLogger(__name__)
        self.sanitizer = DataSanitizer()
        self.collection_name = collection_name

        # Get settings if not provided
        if settings is None:
            try:
                from faultmaven.config.settings import get_settings
                settings = get_settings()
            except Exception:
                settings = None

        # Use injected vector backend or get from factory (Principle 1: Deployment Agnostic)
        if vector_backend is not None:
            self._vector_backend = vector_backend
            self.logger.info(f"Using injected vector backend: {vector_backend.get_backend_type().value}")
        else:
            # Fall back to factory for backwards compatibility
            from faultmaven.infrastructure.vector.factory import get_vector_backend
            self._vector_backend = get_vector_backend()
            self.logger.info(f"Using factory vector backend: {self._vector_backend.get_backend_type().value}")

        # Initialize sentence transformer for embeddings using cached model
        self.embedding_model = model_cache.get_bge_m3_model()
        if self.embedding_model is None:
            self.logger.error("Failed to load BGE-M3 embedding model from cache")
            raise RuntimeError("BGE-M3 model unavailable - knowledge ingestion cannot proceed")
        else:
            self.logger.debug("Using cached BGE-M3 embedding model")

        # Supported file extensions
        self.supported_extensions = {
            ".txt": self._extract_text_txt,
            ".md": self._extract_text_txt,
            ".pdf": self._extract_text_pdf,
            ".docx": self._extract_text_docx,
            ".csv": self._extract_text_csv,
            ".json": self._extract_text_json,
            ".yaml": self._extract_text_yaml,
            ".yml": self._extract_text_yaml,
        }

    @trace("knowledge_base_ingest_document")
    async def ingest_document(
        self,
        file_path: str,
        title: str,
        document_type: str = "troubleshooting_guide",
        tags: Optional[List[str]] = None,
        source_url: Optional[str] = None,
        document_id: Optional[str] = None,
    ) -> str:
        """
        Ingest a document into the knowledge base (background task)

        Args:
            file_path: Path to the document file
            title: Document title
            document_type: Type of document
            tags: Optional tags for categorization
            source_url: Optional source URL
            document_id: Optional document ID to use (generates new if not provided)

        Returns:
            Document ID of the ingested document
        """
        if document_id is None:
            document_id = str(uuid.uuid4())

        try:
            self.logger.info(f"Starting ingestion of document: {title}")

            # Extract text content
            content = await self._extract_text(file_path)
            if not content:
                raise ValueError(f"Could not extract text from {file_path}")

            # Sanitize content
            sanitized_content = self.sanitizer.sanitize(content)

            # Create document object
            document = KnowledgeBaseDocument(
                document_id=document_id,
                title=title,
                content=sanitized_content,
                document_type=document_type,
                tags=tags or [],
                source_url=source_url,
            )

            # Process and store in chunks
            await self._process_and_store(document)

            self.logger.info(f"Successfully ingested document: {title}")
            return document_id

        except Exception as e:
            self.logger.error(f"Failed to ingest document {title}: {e}")
            raise

    async def _extract_text(self, file_path: str) -> str:
        """
        Extract text content from file based on its extension

        Args:
            file_path: Path to the file

        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")

        extractor = self.supported_extensions[file_extension]
        return await extractor(file_path)

    async def _extract_text_txt(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    async def _extract_text_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, "rb") as f:
                pdf_reader = pypdf.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            raise

    async def _extract_text_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise

    async def _extract_text_csv(self, file_path: str) -> str:
        """Extract text from CSV files"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            self.logger.error(f"Failed to extract text from CSV {file_path}: {e}")
            raise

    async def _extract_text_json(self, file_path: str) -> str:
        """Extract text from JSON files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                import json

                data = json.load(f)
                return json.dumps(data, indent=2)
        except Exception as e:
            self.logger.error(f"Failed to extract text from JSON {file_path}: {e}")
            raise

    async def _extract_text_yaml(self, file_path: str) -> str:
        """Extract text from YAML files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                import yaml

                data = yaml.safe_load(f)
                return yaml.dump(data, default_flow_style=False)
        except Exception as e:
            self.logger.error(f"Failed to extract text from YAML {file_path}: {e}")
            raise

    async def _process_and_store(self, document: KnowledgeBaseDocument):
        """
        Process document content and store via IVectorBackend

        Args:
            document: Document to process and store
        """
        # Split content into chunks
        chunks = self._split_content(document.content)

        # Generate embeddings and create VectorDocument objects
        vector_docs = []
        for i, chunk in enumerate(chunks):
            embedding = self.embedding_model.encode(chunk)
            chunk_id = f"{document.document_id}_chunk_{i}"
            metadata = {
                "document_id": document.document_id,
                "title": document.title,
                "document_type": document.document_type,
                "tags": ",".join(document.tags) if document.tags else "",
                "source_url": document.source_url or "",
                "chunk_index": i,
                "total_chunks": len(chunks),
                "created_at": document.created_at.isoformat(),
            }

            vector_docs.append(VectorDocument(
                id=chunk_id,
                content=chunk,
                embedding=embedding.tolist(),
                metadata=metadata,
            ))

        # Store via IVectorBackend interface
        count = await self._vector_backend.upsert(vector_docs, collection=self.collection_name)

        self.logger.info(
            f"Stored {count} chunks for document {document.document_id}"
        )

    def _split_content(
        self, content: str, chunk_size: int = 1000, overlap: int = 200
    ) -> List[str]:
        """
        Split content into overlapping chunks

        Args:
            content: Content to split
            chunk_size: Maximum size of each chunk
            overlap: Overlap between chunks

        Returns:
            List of content chunks
        """
        if len(content) <= chunk_size:
            return [content]

        chunks = []
        start = 0

        while start < len(content):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(content):
                # Look for sentence endings
                for i in range(end, max(start + chunk_size - 100, start), -1):
                    if content[i] in ".!?":
                        end = i + 1
                        break

            chunk = content[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap
            if start >= len(content):
                break

        return chunks

    @trace("knowledge_base_search")
    async def search(
        self,
        query: str,
        n_results: int = 5,
        filter_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Search the knowledge base

        Args:
            query: Search query
            n_results: Number of results to return
            filter_metadata: Optional metadata filters

        Returns:
            List of search results with documents and metadata
        """
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode(query).tolist()

            # Search via IVectorBackend interface
            results: List[VectorSearchResult] = await self._vector_backend.search(
                query_embedding=query_embedding,
                top_k=n_results,
                collection=self.collection_name,
                filter=filter_metadata,
            )

            # Format results for compatibility
            formatted_results = []
            for result in results:
                formatted_results.append({
                    "document": result.content,
                    "metadata": result.metadata,
                    "distance": 1 - result.score,  # Convert score to distance
                    "relevance_score": result.score,
                })

            return formatted_results

        except Exception as e:
            self.logger.error(f"Search failed: {e}")
            return []

    async def delete_document(self, document_id: str) -> bool:
        """
        Delete a document and all its chunks from the knowledge base

        Args:
            document_id: ID of the document to delete

        Returns:
            True if deletion was successful, False if document not found
        """
        try:
            self.logger.info(f"Attempting to delete document {document_id}")

            # Get all chunks for this document to find their IDs
            # First search to find chunk IDs with this document_id
            docs = await self._vector_backend.get(
                ids=[],  # Empty to get all, then filter by metadata
                collection=self.collection_name,
            )

            # Find chunk IDs that belong to this document
            chunk_ids = []
            for doc in docs:
                if doc.metadata and doc.metadata.get("document_id") == document_id:
                    chunk_ids.append(doc.id)

            # If no chunks found by metadata, try pattern-based deletion
            if not chunk_ids:
                # Try to delete using document_id prefix pattern
                # Most backends store chunks as {document_id}_chunk_{i}
                for i in range(100):  # Assume max 100 chunks
                    chunk_ids.append(f"{document_id}_chunk_{i}")

            if chunk_ids:
                deleted = await self._vector_backend.delete(chunk_ids, collection=self.collection_name)
                self.logger.info(f"Deleted {deleted} chunks for document {document_id}")
                return deleted > 0
            else:
                self.logger.warning(f"No chunks found for document {document_id}")
                return False

        except Exception as e:
            self.logger.error(f"Failed to delete document {document_id}: {e}")
            return False

    async def get_collection_stats(self) -> Dict[str, Any]:
        """
        Get statistics about the knowledge base collection

        Returns:
            Dictionary with collection statistics
        """
        try:
            count = await self._vector_backend.count(collection=self.collection_name)

            # Get health info for additional stats
            health = await self._vector_backend.health_check()

            return {
                "total_chunks": count,
                "collection_name": self.collection_name,
                "backend_type": self._vector_backend.get_backend_type().value,
                "backend_health": health.get("status", "unknown"),
            }

        except Exception as e:
            self.logger.error(f"Failed to get collection stats: {e}")
            return {}

    async def ingest_document_object(
        self,
        document: KnowledgeBaseDocument,
    ) -> str:
        """
        Ingest a document object into the knowledge base (for API uploads)

        Args:
            document: KnowledgeBaseDocument object with content already loaded

        Returns:
            Job ID for tracking the ingestion process
        """
        try:
            self.logger.info(f"Starting ingestion of document: {document.title}")

            # Sanitize content (already done in API, but double-check)
            sanitized_content = self.sanitizer.sanitize(document.content)

            # Update document with sanitized content
            document.content = sanitized_content

            # Process and store in chunks
            await self._process_and_store(document)

            self.logger.info(f"Successfully ingested document: {document.title}")

            # Generate job ID for tracking (in a real system, this would be stored)
            job_id = f"job_{document.document_id}"
            return job_id

        except Exception as e:
            self.logger.error(f"Failed to ingest document {document.title}: {e}")
            raise

    async def get_job_status(self, job_id: str) -> Optional[Dict[str, Any]]:
        """
        Get the status of an ingestion job

        Args:
            job_id: Job identifier

        Returns:
            Job status information or None if not found
        """
        # For now, return a simple completed status
        # In a real system, this would track actual job progress
        if job_id.startswith("job_"):
            document_id = job_id.replace("job_", "")
            return {
                "job_id": job_id,
                "document_id": document_id,
                "status": "completed",
                "progress": 100,
                "created_at": "2025-01-01T00:00:00",
                "completed_at": "2025-01-01T00:00:01",
            }
        return None

    async def list_documents(
        self,
        document_type: Optional[str] = None,
        tags: Optional[List[str]] = None,
        limit: int = 50,
        offset: int = 0,
    ) -> List[KnowledgeBaseDocument]:
        """
        List documents in the knowledge base

        Args:
            document_type: Filter by document type
            tags: Filter by tags
            limit: Maximum number of documents
            offset: Number of documents to skip

        Returns:
            List of documents
        """
        try:
            # Build filter for search
            filter_dict = {}
            if document_type:
                filter_dict["document_type"] = document_type
            if tags:
                filter_dict["tags"] = {"$contains": tags[0]}

            # Use search with empty query to list documents
            # Note: This is a simplified implementation - actual listing would
            # need a dedicated list method on the backend
            collections = await self._vector_backend.list_collections()

            # For now, return empty - proper implementation would use backend's list method
            return []

        except Exception as e:
            self.logger.error(f"Failed to list documents: {e}")
            return []

    async def get_document(self, document_id: str) -> Optional[KnowledgeBaseDocument]:
        """
        Get a specific document by ID

        Args:
            document_id: Document identifier

        Returns:
            Document object or None if not found
        """
        try:
            # Get all chunks for this document
            chunk_ids = [f"{document_id}_chunk_{i}" for i in range(100)]
            docs = await self._vector_backend.get(chunk_ids, collection=self.collection_name)

            if not docs:
                return None

            # Filter to only docs that actually exist and belong to this document
            valid_docs = [d for d in docs if d.metadata and d.metadata.get("document_id") == document_id]

            if not valid_docs:
                return None

            # Sort by chunk_index and combine content
            valid_docs.sort(key=lambda d: d.metadata.get("chunk_index", 0))
            content = " ".join(d.content for d in valid_docs)
            metadata = valid_docs[0].metadata

            doc = KnowledgeBaseDocument(
                document_id=document_id,
                title=metadata.get("title", ""),
                content=content,
                document_type=metadata.get("document_type", ""),
                tags=(
                    metadata.get("tags", "").split(",") if metadata.get("tags") else []
                ),
                source_url=metadata.get("source_url"),
            )

            return doc

        except Exception as e:
            self.logger.error(f"Failed to get document {document_id}: {e}")
            return None

    async def search_documents(
        self,
        query: str,
        document_type: Optional[str] = None,
        tags: Optional[List[str]] = None,
        limit: int = 10,
    ) -> List[Dict[str, Any]]:
        """
        Search documents and return results with scores

        Args:
            query: Search query
            document_type: Filter by document type
            tags: Filter by tags
            limit: Maximum number of results

        Returns:
            List of search results with document info and scores
        """
        try:
            # Build filter metadata
            filter_metadata = {}
            if document_type:
                filter_metadata["document_type"] = document_type
            if tags:
                filter_metadata["tags"] = {"$contains": tags[0]}

            # Search using existing search method
            results = await self.search(
                query=query,
                n_results=limit,
                filter_metadata=filter_metadata if filter_metadata else None,
            )

            # Format for API response
            formatted_results = []
            for result in results:
                metadata = result["metadata"]
                formatted_result = {
                    "document_id": metadata.get("document_id"),
                    "title": metadata.get("title"),
                    "document_type": metadata.get("document_type"),
                    "tags": (
                        metadata.get("tags", "").split(",")
                        if metadata.get("tags")
                        else []
                    ),
                    "score": result["relevance_score"],
                    "snippet": (
                        result["document"][:200] + "..."
                        if len(result["document"]) > 200
                        else result["document"]
                    ),
                }
                formatted_results.append(formatted_result)

            return formatted_results

        except Exception as e:
            self.logger.error(f"Search failed: {e}")
            return []
