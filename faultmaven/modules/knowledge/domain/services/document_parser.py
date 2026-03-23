"""Document parser for the conversion pipeline.

Extracts plain text from supported file formats (PDF, DOCX, TXT, Markdown, HTML).
Does not interpret or transform content — that is the LLM's job.
"""

import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentParser:
    """Extracts plain text from uploaded documents."""

    SUPPORTED_FORMATS = {
        "application/pdf": "pdf",
        "application/msword": "docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
        "text/markdown": "markdown",
        "text/html": "html",
    }

    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "txt",
        ".md": "markdown",
        ".markdown": "markdown",
        ".html": "html",
        ".htm": "html",
    }

    def parse(self, file_path: Path, content_type: Optional[str] = None) -> str:
        """Extract plain text from a document file.

        Args:
            file_path: Path to the document file.
            content_type: MIME type hint (used if extension is ambiguous).

        Returns:
            Extracted text content.

        Raises:
            ValueError: If the file format is unsupported or parsing fails.
        """
        fmt = self._detect_format(file_path, content_type)

        extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "txt": self._extract_txt,
            "markdown": self._extract_markdown,
            "html": self._extract_html,
        }

        extractor = extractors.get(fmt)
        if not extractor:
            raise ValueError(f"Unsupported format: {fmt}")

        try:
            text = extractor(file_path)
        except Exception as e:
            raise ValueError(f"Failed to parse {file_path.name}: {e}") from e

        if not text or not text.strip():
            raise ValueError(
                f"No extractable text found in {file_path.name}. "
                "The file may be image-only (scanned PDF) or empty."
            )

        return text

    def _detect_format(self, file_path: Path, content_type: Optional[str]) -> str:
        """Determine format from MIME type or file extension."""
        if content_type and content_type in self.SUPPORTED_FORMATS:
            return self.SUPPORTED_FORMATS[content_type]

        ext = file_path.suffix.lower()
        if ext in self.SUPPORTED_EXTENSIONS:
            return self.SUPPORTED_EXTENSIONS[ext]

        raise ValueError(
            f"Unsupported file type: {ext or content_type}. "
            f"Supported: PDF, DOCX, TXT, Markdown, HTML"
        )

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ValueError("pypdf is required for PDF parsing")

        reader = PdfReader(str(file_path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

        if not pages:
            raise ValueError(
                "No extractable text in PDF. The file may be scanned/image-only."
            )

        return "\n\n".join(pages)

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX preserving structure."""
        try:
            from docx import Document
        except ImportError:
            raise ValueError("python-docx is required for DOCX parsing")

        doc = Document(str(file_path))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Convert heading styles to markdown headings
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if style_name.startswith("heading"):
                    try:
                        level = int(style_name.replace("heading", "").strip())
                        text = "#" * level + " " + text
                    except ValueError:
                        pass

            parts.append(text)

        # Extract tables as markdown tables
        for table in doc.tables:
            table_text = self._docx_table_to_markdown(table)
            if table_text:
                parts.append(table_text)

        return "\n\n".join(parts)

    def _docx_table_to_markdown(self, table) -> str:
        """Convert a DOCX table to markdown format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) < 1:
            return ""

        # Add header separator after first row
        header = rows[0]
        separator = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
        return "\n".join([header, separator] + rows[1:])

    def _read_text_with_fallback(self, file_path: Path) -> str:
        """Read text file with encoding fallback: UTF-8 → Latin-1."""
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            pass
        try:
            return file_path.read_text(encoding="latin-1")
        except UnicodeDecodeError:
            raise ValueError(
                "Cannot decode file — it is not valid UTF-8 or Latin-1 text. "
                "Please re-save the file as UTF-8 and try again."
            )

    def _extract_txt(self, file_path: Path) -> str:
        """Read plain text file."""
        return self._read_text_with_fallback(file_path)

    def _extract_markdown(self, file_path: Path) -> str:
        """Read markdown file, stripping embedded HTML comments."""
        text = self._read_text_with_fallback(file_path)
        # Strip HTML comments
        text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
        return text

    def _extract_html(self, file_path: Path) -> str:
        """Extract text from HTML, preserving structural elements."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ValueError("beautifulsoup4 is required for HTML parsing")

        raw_html = self._read_text_with_fallback(file_path)
        soup = BeautifulSoup(raw_html, "html.parser")

        # Remove non-content elements
        for tag in soup.find_all(
            ["nav", "footer", "aside", "script", "style", "header"]
        ):
            tag.decompose()

        parts = []
        for element in soup.find_all(
            [
                "h1",
                "h2",
                "h3",
                "h4",
                "h5",
                "h6",
                "p",
                "pre",
                "code",
                "ul",
                "ol",
                "table",
                "li",
            ]
        ):
            if element.name and element.name.startswith("h"):
                level = int(element.name[1])
                parts.append("#" * level + " " + element.get_text(strip=True))
            elif element.name in ("pre", "code"):
                code_text = element.get_text()
                if code_text.strip():
                    parts.append(f"```\n{code_text}\n```")
            elif element.name == "li":
                parts.append(f"- {element.get_text(strip=True)}")
            elif element.name == "table":
                parts.append(self._html_table_to_markdown(element))
            elif element.name == "p":
                text = element.get_text(strip=True)
                if text:
                    parts.append(text)

        return "\n\n".join(parts)

    def _html_table_to_markdown(self, table_element) -> str:
        """Convert an HTML table to markdown format."""
        rows = []
        for tr in table_element.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            if cells:
                rows.append("| " + " | ".join(cells) + " |")

        if len(rows) < 1:
            return ""

        header = rows[0]
        num_cols = len(table_element.find("tr").find_all(["td", "th"]))
        separator = "| " + " | ".join(["---"] * num_cols) + " |"
        return "\n".join([header, separator] + rows[1:])
