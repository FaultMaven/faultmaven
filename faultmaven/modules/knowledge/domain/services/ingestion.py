"""ingestion.py

Purpose: Knowledge base population pipeline

Requirements:
--------------------------------------------------------------------------------
• Handle background processing
• Support multiple file formats
• Generate embeddings and store in ChromaDB

Key Components:
--------------------------------------------------------------------------------
  class KnowledgeIngester: ...
  @background_task def ingest_document(...)

Technology Stack:
--------------------------------------------------------------------------------
PyPDF2, python-docx, sentence-transformers

Core Design Principles:
--------------------------------------------------------------------------------
• Privacy-First: Sanitize all external-bound data
• Resilience: Implement retries and fallbacks
• Cost-Efficiency: Use semantic caching
• Extensibility: Use interfaces for pluggable components
• Observability: Add tracing spans for key operations
"""

import logging
import os
import re
import threading
import uuid
from datetime import datetime, timezone
from typing import Any, Callable, Dict, List, Optional

import chromadb
import pandas as pd
import pypdf
from chromadb.config import Settings
from docx import Document

from faultmaven.infrastructure.model_cache import model_cache
from faultmaven.infrastructure.observability.tracing import trace
from faultmaven.infrastructure.security.redaction import DataSanitizer
from faultmaven.models import KnowledgeBaseDocument


def _call_with_timeout(fn: Callable[[], Any], timeout_s: float, what: str) -> Any:
    """Run ``fn()`` on a daemon thread, raising ``TimeoutError`` if it has not
    returned within ``timeout_s`` seconds.

    ChromaDB's ``HttpClient`` performs blocking network round-trips with no
    caller-facing timeout, so an unreachable server would otherwise hang
    application startup until a readiness probe gives up. The worker thread is
    a daemon: on timeout we abandon it (it dies with the process or when the
    underlying socket finally errors) and return control to the caller.
    """
    box: Dict[str, Any] = {}

    def _run() -> None:
        try:
            box["value"] = fn()
        except BaseException as exc:  # noqa: BLE001 - re-raised to caller below
            box["error"] = exc

    worker = threading.Thread(target=_run, name="chromadb-connect", daemon=True)
    worker.start()
    worker.join(timeout_s)
    if worker.is_alive():
        raise TimeoutError(f"{what} did not respond within {timeout_s:.0f}s")
    if "error" in box:
        raise box["error"]
    return box.get("value")


class KnowledgeIngester:
    """Handles asynchronous ingestion of documents into the knowledge base"""

    def __init__(
        self, chroma_persist_directory: str = "./data/chroma-kb", settings=None
    ):
        self.logger = logging.getLogger(__name__)
        self.sanitizer = DataSanitizer()

        # Get settings if not provided
        if settings is None:
            try:
                from faultmaven.config.settings import get_settings

                settings = get_settings()
            except Exception:
                settings = None

        # Initialize ChromaDB - default to K8s cluster for production-like development
        if settings:
            # Use settings-based configuration
            chromadb_url = settings.database.chromadb_url
            chromadb_host = settings.database.chromadb_host
            chromadb_port = settings.database.chromadb_port
            chromadb_auth_token = (
                settings.database.chromadb_auth_token.get_secret_value()
                if settings.database.chromadb_auth_token
                else "faultmaven-dev-chromadb-2025"
            )
        else:
            # No fallback - unified settings system is mandatory
            from faultmaven.models.exceptions import KnowledgeBaseError

            raise KnowledgeBaseError(
                "Knowledge ingestion requires unified settings system to be available",
                error_code="KNOWLEDGE_CONFIG_ERROR",
                context={"settings_available": settings is not None},
            )

        # Degraded state: when an external ChromaDB is configured but
        # unreachable, we do NOT block startup. self._collection stays None and
        # the `collection` property raises a typed error on use; callers (and
        # the KB bootstrap) treat that as "KB temporarily unavailable".
        self.degraded = False
        self.chroma_client = None
        self._collection = None
        connect_timeout = float(
            getattr(settings.database, "chromadb_connect_timeout", 5.0)
        )

        # Single unified collection for all KB scopes (global, team, personal).
        # Must match UnifiedKBConfig.get_collection_name() = "faultmaven_kb".
        # Scope isolation via metadata filtering, not separate collections.
        def _open_http(label: str, **client_kwargs: Any) -> None:
            self.logger.info(f"Using ChromaDB HTTP client at {label}")

            def _build():
                client = chromadb.HttpClient(**client_kwargs)
                # Force a real round-trip here so an unreachable / mis-auth'd
                # server fails inside the timeout guard, not on first query.
                collection = client.get_or_create_collection(
                    name="faultmaven_kb",
                    metadata={"description": "FaultMaven Knowledge Base"},
                )
                return client, collection

            try:
                self.chroma_client, self._collection = _call_with_timeout(
                    _build, connect_timeout, f"ChromaDB at {label}"
                )
            except Exception as exc:  # unreachable, timeout, auth failure
                self.degraded = True
                self.logger.error(
                    "ChromaDB unavailable at %s (%r) — knowledge ingestion and "
                    "search are disabled until it recovers; startup continues.",
                    label,
                    exc,
                )

        if chromadb_url:
            host = (
                chromadb_url.replace("http://", "")
                .replace("https://", "")
                .split(":")[0]
            )
            _open_http(
                chromadb_url,
                host=host,
                port=int(chromadb_url.split(":")[-1]),
                settings=Settings(anonymized_telemetry=False, allow_reset=True),
            )
        elif chromadb_host != "localhost":
            # K8s cluster or external HTTP client
            _open_http(
                f"{chromadb_host}:{chromadb_port}",
                host=chromadb_host,
                port=chromadb_port,
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True,
                    chroma_client_auth_provider="chromadb.auth.token_authn.TokenAuthClientProvider",
                    chroma_client_auth_credentials=chromadb_auth_token,
                ),
            )
        else:
            # Local development with persistent client (no network — cannot hang)
            kb_dir = getattr(
                settings.database, "chromadb_kb_persist_dir", chroma_persist_directory
            )
            self.logger.info(f"Using ChromaDB PersistentClient at {kb_dir}")
            self.chroma_client = chromadb.PersistentClient(
                path=kb_dir,
                settings=Settings(anonymized_telemetry=False, allow_reset=True),
            )
            self._collection = self.chroma_client.get_or_create_collection(
                name="faultmaven_kb",
                metadata={"description": "FaultMaven Knowledge Base"},
            )

        # Fail fast at startup if BGE-M3 can't load (also warms the lazy cache).
        # Embedding itself goes through model_cache.aembed_* at the call sites
        # below, which run the encode off the event loop.
        if model_cache.get_bge_m3_model() is None:
            self.logger.error("Failed to load BGE-M3 embedding model from cache")
            raise RuntimeError(
                "BGE-M3 model unavailable - knowledge ingestion cannot proceed"
            )
        self.logger.debug("Cached BGE-M3 embedding model available")

        # Supported file extensions
        self.supported_extensions = {
            ".txt": self._extract_text_txt,
            ".md": self._extract_text_txt,
            ".pdf": self._extract_text_pdf,
            ".docx": self._extract_text_docx,
            ".csv": self._extract_text_csv,
            ".json": self._extract_text_json,
            ".yaml": self._extract_text_yaml,
            ".yml": self._extract_text_yaml,
        }

    @property
    def collection(self):
        """The ChromaDB collection, or a typed error if the store is degraded.

        When an external ChromaDB was configured but unreachable at startup,
        ``_collection`` is None and any access raises ``KnowledgeBaseError``.
        Existing per-method try/except blocks turn that into a graceful
        "KB unavailable" outcome instead of a crash or hang.
        """
        if self._collection is None:
            from faultmaven.models.exceptions import KnowledgeBaseError

            raise KnowledgeBaseError(
                "Knowledge base vector store is unavailable (ChromaDB not reachable)",
                error_code="KNOWLEDGE_STORE_UNAVAILABLE",
            )
        return self._collection

    @trace("knowledge_base_ingest_document")
    async def ingest_document(
        self,
        file_path: str,
        title: str,
        document_type: str = "troubleshooting_guide",
        tags: Optional[List[str]] = None,
        source_url: Optional[str] = None,
        document_id: Optional[str] = None,
        scope: str = "global",
        owner_id: Optional[str] = None,
    ) -> str:
        """
        Ingest a document into the knowledge base (background task)

        Args:
            file_path: Path to the document file
            title: Document title
            document_type: Type of document
            tags: Optional tags for categorization
            source_url: Optional source URL
            document_id: Optional document ID to use (generates new if not provided)
            scope: Visibility scope (global, personal, team). Team visibility is
                carried by the share table (resource_shares), not here — this
                path only tags the immutable metadata floor (owner/global).
            owner_id: Owner user ID (required for personal scope)

        Returns:
            Document ID of the ingested document
        """
        if document_id is None:
            document_id = str(uuid.uuid4())

        try:
            self.logger.info(f"Starting ingestion of document: {title}")

            # Extract text content
            content = await self._extract_text(file_path)
            if not content:
                raise ValueError(f"Could not extract text from {file_path}")

            # Sanitize content
            sanitized_content = await self.sanitizer.asanitize(content)

            # Create document object
            now = datetime.now(timezone.utc).isoformat()
            document = KnowledgeBaseDocument(
                document_id=document_id,
                title=title,
                content=sanitized_content,
                document_type=document_type,
                tags=tags or [],
                source_url=source_url,
                scope=scope,
                owner_id=owner_id,
                created_at=now,
                updated_at=now,
            )

            # Process and store in chunks
            await self._process_and_store(document)

            self.logger.info(f"Successfully ingested document: {title}")
            return document_id

        except Exception as e:
            self.logger.error(f"Failed to ingest document {title}: {e}")
            raise

    async def _extract_text(self, file_path: str) -> str:
        """
        Extract text content from file based on its extension

        Args:
            file_path: Path to the file

        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")

        extractor = self.supported_extensions[file_extension]
        return await extractor(file_path)

    async def _extract_text_txt(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    async def _extract_text_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, "rb") as f:
                pdf_reader = pypdf.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            raise

    async def _extract_text_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise

    async def _extract_text_csv(self, file_path: str) -> str:
        """Extract text from CSV files"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            self.logger.error(f"Failed to extract text from CSV {file_path}: {e}")
            raise

    async def _extract_text_json(self, file_path: str) -> str:
        """Extract text from JSON files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                import json

                data = json.load(f)
                return json.dumps(data, indent=2)
        except Exception as e:
            self.logger.error(f"Failed to extract text from JSON {file_path}: {e}")
            raise

    async def _extract_text_yaml(self, file_path: str) -> str:
        """Extract text from YAML files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                import yaml

                data = yaml.safe_load(f)
                return yaml.dump(data, default_flow_style=False)
        except Exception as e:
            self.logger.error(f"Failed to extract text from YAML {file_path}: {e}")
            raise

    async def _process_and_store(self, document: KnowledgeBaseDocument):
        """
        Process document content and store in ChromaDB

        Args:
            document: Document to process and store
        """
        # Extract frontmatter metadata for RAG enrichment (domain, service, etc.)
        frontmatter_meta = self._extract_frontmatter_metadata(document.content)

        # Split content into chunks
        chunks = self._split_content(document.content)
        if not chunks:
            # Empty / whitespace-only document → nothing to index. Skip before
            # embedding: aembed_texts([]) returns [] (not None), and
            # collection.add(embeddings=[], ...) errors on an empty add. Mirrors
            # the guards in vector_storage / knowledge_service.
            self.logger.warning(
                f"No chunks produced for document {document.document_id}, "
                "skipping vector indexing"
            )
            return

        # Batch-embed all chunks off the event loop via the model_cache async
        # boundary (availability is guaranteed by the startup check).
        embeddings = await model_cache.aembed_texts(chunks)
        if embeddings is None:
            raise RuntimeError("BGE-M3 model unavailable during ingestion")

        # Prepare metadata for each chunk
        metadatas = []
        ids = []

        for i, chunk in enumerate(chunks):
            chunk_id = f"{document.document_id}_chunk_{i}"
            # Metadata scope carries only the immutable floor: 'global' (platform)
            # vs 'personal' (owner-only). A team item is tagged 'personal' here —
            # its team visibility lives in the share table and is resolved into an
            # id allowlist at query time. Writing 'team'/team_id here would orphan
            # the chunk on unshare, and writing 'global' would leak it to everyone
            # (ADR-013 §D4 / ADR-011 D3).
            _raw_scope = getattr(document, "scope", None) or "global"
            _meta_scope = "global" if _raw_scope == "global" else "personal"
            metadata = {
                "document_id": document.document_id,
                "parent_document_id": document.document_id,
                "title": document.title,
                "document_type": document.document_type,
                "tags": ",".join(document.tags) if document.tags else "",
                "source_url": document.source_url or "",
                "scope": _meta_scope,
                "owner_id": getattr(document, "owner_id", None) or "",
                "chunk_index": i,
                "total_chunks": len(chunks),
                "created_at": document.created_at,
            }

            # Enrich with frontmatter fields for hybrid search and staleness
            if frontmatter_meta.get("domain"):
                metadata["domain"] = frontmatter_meta["domain"]
            if frontmatter_meta.get("service"):
                metadata["service"] = frontmatter_meta["service"]
            if frontmatter_meta.get("last_updated"):
                metadata["last_updated"] = frontmatter_meta["last_updated"]
            if frontmatter_meta.get("status"):
                metadata["status"] = frontmatter_meta["status"]
            if frontmatter_meta.get("severity"):
                metadata["severity"] = frontmatter_meta["severity"]
            if frontmatter_meta.get("symptom_class"):
                metadata["symptom_class"] = frontmatter_meta["symptom_class"]

            ids.append(chunk_id)
            metadatas.append(metadata)

        # Store in ChromaDB
        self.collection.add(
            embeddings=embeddings, documents=chunks, metadatas=metadatas, ids=ids
        )

        self.logger.info(
            f"Stored {len(chunks)} chunks for document {document.document_id}"
        )

    @staticmethod
    def _extract_frontmatter_metadata(content: str) -> Dict[str, str]:
        """Extract RAG-relevant metadata from YAML frontmatter."""
        from faultmaven.utils.frontmatter import extract_frontmatter_metadata

        return extract_frontmatter_metadata(content)

    # Maximum chunk size — sections exceeding this get split at sentence boundaries
    MAX_CHUNK_CHARS = 3000
    # Minimum chunk size — tiny sections get merged with the next section
    MIN_CHUNK_CHARS = 100

    def _split_content(
        self, content: str, chunk_size: int = 1000, overlap: int = 200
    ) -> List[str]:
        """Structure-aware content splitting for runbooks and documentation.

        Splits on document structure boundaries (markdown headers, numbered
        steps, horizontal rules) rather than fixed character counts. This
        preserves semantic units — a diagnostic step stays with its
        conditional, a remediation procedure stays intact.

        Fallback: If no structure is detected (plain text without headers),
        falls back to sentence-boundary splitting at MAX_CHUNK_CHARS.

        Variable chunk sizes are intentional — a 200-char config parameter
        description is one chunk, a 2500-char procedure section is one chunk.
        The embedding model handles this fine.
        """
        # Strip frontmatter before chunking
        stripped = re.sub(
            r"^---\s*\n.*?\n---\s*\n", "", content, count=1, flags=re.DOTALL
        )
        stripped = stripped.strip()

        if not stripped:
            return [content.strip()] if content.strip() else []

        # Try structure-aware splitting first
        sections = self._split_by_structure(stripped)

        if len(sections) <= 1 and len(stripped) > self.MAX_CHUNK_CHARS:
            # No structure detected — fall back to sentence-boundary splitting
            sections = self._split_by_sentences(stripped)

        # Post-process: merge tiny sections, split oversized ones
        return self._normalize_chunks(sections)

    @staticmethod
    def _split_by_structure(content: str) -> List[str]:
        """Split content at markdown structural boundaries.

        Recognizes:
          - Markdown headers (# through ####)
          - Numbered steps (1. 2. 3.)
          - Horizontal rules (--- or ***)
          - Blank-line-separated blocks of text
        """
        # Split on markdown headers (##, ###, ####) — keep the header with its section
        # Pattern: one or more blank lines followed by a header line
        header_pattern = re.compile(r"\n(?=#{1,4}\s+\S)")

        parts = header_pattern.split(content)
        sections = [p.strip() for p in parts if p.strip()]

        # If we got meaningful splits, return them
        if len(sections) > 1:
            return sections

        # No headers found — try splitting on horizontal rules
        hr_pattern = re.compile(r"\n\s*(?:---+|\*\*\*+|___+)\s*\n")
        parts = hr_pattern.split(content)
        sections = [p.strip() for p in parts if p.strip()]
        if len(sections) > 1:
            return sections

        # No rules either — return as single section
        return [content.strip()]

    @staticmethod
    def _split_by_sentences(content: str, max_size: int = 3000) -> List[str]:
        """Fallback: split at sentence boundaries with max size limit."""
        if len(content) <= max_size:
            return [content]

        chunks = []
        current = []
        current_len = 0

        for line in content.split("\n"):
            line_len = len(line) + 1  # +1 for newline
            if current_len + line_len > max_size and current:
                chunks.append("\n".join(current))
                current = [line]
                current_len = line_len
            else:
                current.append(line)
                current_len += line_len

        if current:
            chunks.append("\n".join(current))

        return chunks

    def _normalize_chunks(self, sections: List[str]) -> List[str]:
        """Merge tiny sections and split oversized ones."""
        normalized = []
        pending = ""

        for section in sections:
            if pending:
                combined = pending + "\n\n" + section
                if len(combined) <= self.MAX_CHUNK_CHARS:
                    pending = combined
                    continue
                else:
                    # Pending is big enough on its own
                    normalized.append(pending.strip())
                    pending = ""

            if len(section) < self.MIN_CHUNK_CHARS:
                pending = section
            elif len(section) > self.MAX_CHUNK_CHARS:
                # Oversized section — split at sentence boundaries
                sub_chunks = self._split_by_sentences(section, self.MAX_CHUNK_CHARS)
                normalized.extend(sub_chunks)
            else:
                normalized.append(section)

        if pending:
            if normalized:
                # Merge trailing tiny chunk into last chunk
                last = normalized[-1]
                if len(last) + len(pending) + 2 <= self.MAX_CHUNK_CHARS:
                    normalized[-1] = last + "\n\n" + pending
                else:
                    normalized.append(pending.strip())
            else:
                normalized.append(pending.strip())

        return [c for c in normalized if c.strip()]

    @trace("knowledge_base_search")
    async def search(
        self,
        query: str,
        n_results: int = 5,
        filter_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Search the knowledge base

        Args:
            query: Search query
            n_results: Number of results to return
            filter_metadata: Optional metadata filters

        Returns:
            List of search results with documents and metadata
        """
        try:
            # Generate query embedding off the event loop via the async boundary.
            query_embedding = await model_cache.aembed_query(query)
            if query_embedding is None:
                self.logger.error("BGE-M3 unavailable for query embedding")
                return []

            # Prepare where clause for filtering
            where_clause = None
            if filter_metadata:
                where_clause = filter_metadata

            # Search in ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results,
                where=where_clause,
                include=["documents", "metadatas", "distances"],
            )

            # Format results
            formatted_results = []
            if results["documents"] and results["documents"][0]:
                for i, doc in enumerate(results["documents"][0]):
                    result = {
                        "document": doc,
                        "metadata": results["metadatas"][0][i],
                        "distance": results["distances"][0][i],
                        "relevance_score": 1
                        - results["distances"][0][i],  # Convert distance to relevance
                    }
                    formatted_results.append(result)

            return formatted_results

        except Exception as e:
            self.logger.error(f"Search failed: {e}")
            return []

    async def delete_document(self, document_id: str) -> bool:
        """
        Delete a document and all its chunks from the knowledge base

        Args:
            document_id: ID of the document to delete

        Returns:
            True if deletion was successful, False if document not found
        """
        try:
            # Find all chunks for this document
            self.logger.info(f"Attempting to delete document {document_id}")

            # First, let's see what's in the collection
            all_results = self.collection.get(include=["metadatas"], limit=10)
            self.logger.info(
                f"Sample collection contents: {len(all_results.get('metadatas', []))} items"
            )
            if all_results.get("metadatas"):
                for i, meta in enumerate(all_results["metadatas"][:3]):
                    self.logger.info(
                        f"Sample item {i}: {meta.get('document_id', 'no_id')}"
                    )

            results = self.collection.get(
                where={"document_id": document_id}, include=["metadatas"]
            )

            # ChromaDB returns IDs by default in results
            chunk_ids = results.get("ids", [])
            self.logger.info(
                f"Query results for {document_id}: found {len(chunk_ids)} chunk IDs"
            )

            if chunk_ids and len(chunk_ids) > 0:
                # Delete all chunks
                self.collection.delete(ids=chunk_ids)
                self.logger.info(
                    f"Deleted {len(chunk_ids)} chunks for document {document_id}"
                )
                return True
            else:
                self.logger.warning(f"No chunks found for document {document_id}")
                return False

        except Exception as e:
            self.logger.error(f"Failed to delete document {document_id}: {e}")
            return False

    def get_collection_stats(self) -> Dict[str, Any]:
        """
        Get statistics about the knowledge base collection

        Returns:
            Dictionary with collection statistics
        """
        try:
            count = self.collection.count()

            # Get sample of documents to analyze
            sample = self.collection.get(limit=1000, include=["metadatas"])

            # Analyze document types
            doc_types = {}
            tags = {}

            if sample["metadatas"]:
                for metadata in sample["metadatas"]:
                    doc_type = metadata.get("document_type", "unknown")
                    doc_types[doc_type] = doc_types.get(doc_type, 0) + 1

                    tag_list = metadata.get("tags", "").split(",")
                    for tag in tag_list:
                        tag = tag.strip()
                        if tag:
                            tags[tag] = tags.get(tag, 0) + 1

            return {
                "total_chunks": count,
                "document_types": doc_types,
                "top_tags": dict(
                    sorted(tags.items(), key=lambda x: x[1], reverse=True)[:10]
                ),
                "collection_name": self.collection.name,
            }

        except Exception as e:
            self.logger.error(f"Failed to get collection stats: {e}")
            return {}

    async def ingest_document_object(
        self,
        document: KnowledgeBaseDocument,
    ) -> str:
        """
        Ingest a document object into the knowledge base (for API uploads)

        Args:
            document: KnowledgeBaseDocument object with content already loaded

        Returns:
            Job ID for tracking the ingestion process
        """
        try:
            self.logger.info(f"Starting ingestion of document: {document.title}")

            # Sanitize content (already done in API, but double-check)
            sanitized_content = await self.sanitizer.asanitize(document.content)

            # Update document with sanitized content
            document.content = sanitized_content

            # Process and store in chunks
            await self._process_and_store(document)

            self.logger.info(f"Successfully ingested document: {document.title}")

            # Generate job ID for tracking (in a real system, this would be stored)
            job_id = f"job_{document.document_id}"
            return job_id

        except Exception as e:
            self.logger.error(f"Failed to ingest document {document.title}: {e}")
            raise

    async def get_job_status(self, job_id: str) -> Optional[Dict[str, Any]]:
        """
        Get the status of an ingestion job

        Args:
            job_id: Job identifier

        Returns:
            Job status information or None if not found
        """
        # For now, return a simple completed status
        # In a real system, this would track actual job progress
        if job_id.startswith("job_"):
            document_id = job_id.replace("job_", "")
            return {
                "job_id": job_id,
                "document_id": document_id,
                "status": "completed",
                "progress": 100,
                "created_at": "2025-01-01T00:00:00",
                "completed_at": "2025-01-01T00:00:01",
            }
        return None

    async def list_documents(
        self,
        document_type: Optional[str] = None,
        tags: Optional[List[str]] = None,
        limit: int = 50,
        offset: int = 0,
    ) -> List[KnowledgeBaseDocument]:
        """
        List documents in the knowledge base

        Args:
            document_type: Filter by document type
            tags: Filter by tags
            limit: Maximum number of documents
            offset: Number of documents to skip

        Returns:
            List of documents
        """
        try:
            # Build where clause for filtering
            where_clause = {}
            if document_type:
                where_clause["document_type"] = document_type
            if tags:
                # For simplicity, just filter by first tag
                # In a real implementation, would need more complex tag filtering
                where_clause["tags"] = {"$contains": tags[0]}

            # Get documents from ChromaDB
            results = self.collection.get(
                where=where_clause if where_clause else None,
                include=["metadatas"],
                limit=limit,
                offset=offset,
            )

            # Convert to document objects
            documents = []
            seen_doc_ids = set()

            if results["metadatas"]:
                for metadata in results["metadatas"]:
                    doc_id = metadata.get("document_id")
                    if doc_id and doc_id not in seen_doc_ids:
                        seen_doc_ids.add(doc_id)
                        # Create document object from metadata
                        doc = KnowledgeBaseDocument(
                            document_id=doc_id,
                            title=metadata.get("title", ""),
                            content="",  # Don't include full content in list
                            document_type=metadata.get("document_type", ""),
                            tags=(
                                metadata.get("tags", "").split(",")
                                if metadata.get("tags")
                                else []
                            ),
                            source_url=metadata.get("source_url"),
                        )
                        documents.append(doc)

            return documents

        except Exception as e:
            self.logger.error(f"Failed to list documents: {e}")
            return []

    async def get_document(self, document_id: str) -> Optional[KnowledgeBaseDocument]:
        """
        Get a specific document by ID

        Args:
            document_id: Document identifier

        Returns:
            Document object or None if not found
        """
        try:
            # Get all chunks for this document
            results = self.collection.get(
                where={"document_id": document_id}, include=["documents", "metadatas"]
            )

            if not results["documents"] or not results["documents"]:
                return None

            # Reconstruct document from chunks
            chunks = results["documents"]
            metadata = results["metadatas"][0] if results["metadatas"] else {}

            # Combine all chunks to reconstruct content
            content = " ".join(chunks)

            doc = KnowledgeBaseDocument(
                document_id=document_id,
                title=metadata.get("title", ""),
                content=content,
                document_type=metadata.get("document_type", ""),
                tags=(
                    metadata.get("tags", "").split(",") if metadata.get("tags") else []
                ),
                source_url=metadata.get("source_url"),
            )

            return doc

        except Exception as e:
            self.logger.error(f"Failed to get document {document_id}: {e}")
            return None

    async def search_documents(
        self,
        query: str,
        document_type: Optional[str] = None,
        tags: Optional[List[str]] = None,
        limit: int = 10,
    ) -> List[Dict[str, Any]]:
        """
        Search documents and return results with scores

        Args:
            query: Search query
            document_type: Filter by document type
            tags: Filter by tags
            limit: Maximum number of results

        Returns:
            List of search results with document info and scores
        """
        try:
            # Build filter metadata
            filter_metadata = {}
            if document_type:
                filter_metadata["document_type"] = document_type
            if tags:
                filter_metadata["tags"] = {"$contains": tags[0]}

            # Search using existing search method
            results = await self.search(
                query=query,
                n_results=limit,
                filter_metadata=filter_metadata if filter_metadata else None,
            )

            # Format for API response
            formatted_results = []
            for result in results:
                metadata = result["metadata"]
                formatted_result = {
                    "document_id": metadata.get("document_id"),
                    "title": metadata.get("title"),
                    "document_type": metadata.get("document_type"),
                    "tags": (
                        metadata.get("tags", "").split(",")
                        if metadata.get("tags")
                        else []
                    ),
                    "score": result["relevance_score"],
                    "snippet": (
                        result["document"][:200] + "..."
                        if len(result["document"]) > 200
                        else result["document"]
                    ),
                }
                formatted_results.append(formatted_result)

            return formatted_results

        except Exception as e:
            self.logger.error(f"Search failed: {e}")
            return []
